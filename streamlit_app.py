#!/usr/bin/env python3
"""
Bowen Theory RAG — Streamlit web app
Same backend as the desktop app; designed for Railway deployment.
"""

import json
import logging
import os
import re
from pathlib import Path

import citations

# Load .env so local runs pick up API keys without manual export
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Suppress verbose startup logs from PyTorch / transformers / sentence-transformers
import warnings
warnings.filterwarnings("ignore")
for _noisy in ("transformers", "sentence_transformers", "torch", "filelock",
               "urllib3", "huggingface_hub", "tqdm"):
    logging.getLogger(_noisy).setLevel(logging.ERROR)
logging.getLogger("root").setLevel(logging.WARNING)

import numpy as np
import streamlit as st
from scipy import sparse as sp_sparse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    from sentence_transformers import SentenceTransformer
    EMBEDDING_AVAILABLE = True
except ImportError:
    EMBEDDING_AVAILABLE = False

try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except ImportError:
    BM25_AVAILABLE = False

# ── Paths ──────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
REFS_DIR = BASE_DIR / "rag-document-search" / "references"

# ── Authority tiers ─────────────────────────────────────────────────────────
_AUTHORITY_TIERS_DEFAULT = [
    ("Family Therapy in_Clinical_Practice_Chapter", 3.0),
    ("Family Evaluation",                            3.0),
    ("Bowen Basic Series Tape",                      3.0),
    ("BOWEN-KERR INTERVIEW SERIES",                  3.0),
    ("Bowen Family Systems Theory",                  3.0),
    ("Bowen on Triangles",                           3.0),
    ("Bowen Theory and Therapy",                     3.0),
    ("Chronic Anxiety and Defining",                 3.0),
    ("Cancer and the Emotional System",              3.0),
    ("Family and Society Kerr",                      3.0),
    ("Family as a System Kerr",                      3.0),
    ("Family Systems and Therapy Kerr",              3.0),
    ("Physical Illness as the Family Emotional",     3.0),
    ("Psychotherapy Past Present Future",            3.0),
    ("FSJ ",                                         1.3),
    ("Copy of ",                                     1.3),
    ("Family Center Reports",                        1.3),
    ("Papero",                                       1.15),
    ("Friedman",                                     1.15),
    ("Fogarty",                                      1.15),
    ("Guerin",                                       1.15),
    ("Toman",                                        1.15),
]


def _load_authority_tiers() -> list:
    config = BASE_DIR / "authority_tiers.yml"
    if not config.exists():
        return _AUTHORITY_TIERS_DEFAULT
    try:
        import yaml
        with open(config, encoding="utf-8") as f:
            data = yaml.safe_load(f)
        return [(t["pattern"], float(t["multiplier"]))
                for t in data.get("tiers", []) if "pattern" in t]
    except Exception:
        return _AUTHORITY_TIERS_DEFAULT


AUTHORITY_TIERS = _load_authority_tiers()


def authority_boost(doc_name: str) -> float:
    dn = doc_name.lower()
    for pattern, mult in AUTHORITY_TIERS:
        if pattern.lower() in dn:
            return mult
    return 1.0


def _load_author_map() -> list:
    config = BASE_DIR / "author_map.yml"
    if not config.exists():
        return []
    try:
        import yaml
        with open(config, encoding="utf-8") as f:
            data = yaml.safe_load(f)
        return [(a["pattern"], a["author"])
                for a in data.get("authors", []) if "pattern" in a and "author" in a]
    except Exception:
        return []


AUTHOR_MAP = _load_author_map()

# Bibliographic records for citation formatting (sources.yml); [] if absent.
SOURCES = citations.load_sources(BASE_DIR)


def doc_author(doc_name: str) -> str:
    dn = doc_name.lower()
    for pattern, author in AUTHOR_MAP:
        if pattern.lower() in dn:
            return author
    return "Unknown"


def all_known_authors() -> list:
    seen: list = []
    for _, author in AUTHOR_MAP:
        if author not in seen:
            seen.append(author)
    return seen


# ── System prompt ────────────────────────────────────────────────────────────
SYSTEM_PROMPT = (
    "You are a research assistant helping analyse the Bowen Family Systems Theory literature. "
    "STRICT RULES — you must follow these without exception:\n"
    "1. Use ONLY the source excerpts provided in the user message. "
    "Do not draw on any prior training knowledge, general knowledge, or outside information.\n"
    "2. Do not infer, extrapolate, or fill gaps with assumptions. "
    "If the provided excerpts do not address something, say so explicitly rather than guessing.\n"
    "3. Every claim or statement in your response must be directly traceable to a specific excerpt. "
    "Cite the source document in brackets immediately after the claim, e.g. [Document Name].\n"
    "4. If sources conflict or are ambiguous, note the conflict and quote both — do not resolve it yourself.\n"
    "5. Do not add introductory or concluding remarks that go beyond what the sources say.\n"
    "6. If asked about something not covered in the provided excerpts, respond: "
    "'The provided sources do not contain information on this point.'\n"
    "7. Terminology precision — CRITICAL for Bowen theory: Before answering, check whether the "
    "exact term or phrase from the user's query appears verbatim in the provided excerpts. "
    "If it does NOT appear verbatim, you MUST begin your response with a clearly visible notice:\n"
    "⚠️ Note: The exact term/phrase \"[user's term]\" does not appear in the provided sources. "
    "The following is drawn from semantically related content and represents an inference, "
    "not a direct match. Bowen theory uses precise terminology — the absence of a term is "
    "significant and should not be papered over by substituting similar-sounding concepts.\n"
    "This rule exists because, for example, Bowen uses 'emotional contact' with a specific "
    "clinical meaning that is distinct from everyday phrases like 'emotional connection' — "
    "treating them as equivalent would misrepresent the theory."
)

INTERVIEW_SYSTEM_PROMPT = (
    SYSTEM_PROMPT + "\n\n"
    "INTERACTION MODE: INTERVIEW\n"
    "You are a researcher conducting a Bowen Family Systems Theory interview. Your purpose is to "
    "explore how the person is functioning in three specific areas, drawn from the retrieved source "
    "excerpts: (1) defining a self — the ability to hold one's own values, beliefs, and positions "
    "under relationship pressure without either capitulating or becoming reactive; "
    "(2) managing emotional reactivity — recognising and moderating automatic emotional responses "
    "that are driven by anxiety in the system rather than thoughtful choice; "
    "(3) maintaining emotional contact — staying in genuine, non-fused connection with important "
    "people in the system rather than withdrawing into distance or cutoff.\n\n"
    "Rules for this mode:\n"
    "- Ask ONE focused question per turn. Never ask multiple questions at once.\n"
    "- Follow the person's lead — let their answer determine which area to explore next. "
    "Do not mechanically work through a checklist.\n"
    "- After the person responds, offer one brief neutral observation (1 sentence) that names "
    "what you heard without evaluation, then ask your next question.\n"
    "- Your questions should be curious and open — aimed at understanding, not at leading the "
    "person toward a particular answer.\n"
    "- Do not interpret, diagnose, or evaluate the person's functioning. You are gathering "
    "information, not rendering a verdict.\n"
    "- Ground any theoretical framing strictly in the retrieved source excerpts and cite the source.\n"
    "- If the person asks a direct question about theory, answer briefly from the sources "
    "and return to the interview.\n"
    "- Keep your turn to one observation + one question. Do not lecture."
)

COACH_SYSTEM_PROMPT = (
    SYSTEM_PROMPT + "\n\n"
    "INTERACTION MODE: COACHING\n"
    "You are a Bowen Family Systems Theory coach. Your sole purpose is to help the person "
    "think more clearly about their own part in their relationship system — not to advise, "
    "direct, fix, or tell them what to do.\n\n"
    "Rules for this mode:\n"
    "- NEVER tell the person what they should do, ought to do, or could try. "
    "That is not coaching in the Bowen sense — it is advice, and it keeps the person focused "
    "on others or on actions rather than on their own thinking and functioning.\n"
    "- Every question you ask should return the person's attention to their OWN PART: "
    "their own reactions, their own positions, their own patterns, their own choices. "
    "Example orientations: 'What do you notice about your own part in that?' "
    "'How do you tend to respond when that happens?' "
    "'What would it mean for you personally to take a different position there?'\n"
    "- After the person speaks, reflect back in one sentence what you heard about their "
    "functioning (not about the other people they mentioned), then ask ONE question.\n"
    "- If the person asks what they should do, redirect: acknowledge the question, then ask "
    "what they themselves think, or what they notice about their own reactivity to the situation.\n"
    "- When a Bowen concept from the sources is directly relevant, name it briefly and cite it — "
    "then ask how the person sees it in their own behaviour, not in others'.\n"
    "- Keep turns concise: one reflection on their functioning + one question. No lists, no plans."
)

QUIZ_SYSTEM_PROMPT = (
    SYSTEM_PROMPT + "\n\n"
    "INTERACTION MODE: QUIZ\n"
    "You are quizzing the user on Bowen Family Systems Theory, drawing exclusively from the "
    "retrieved source excerpts. Rules for this mode:\n"
    "- Ask ONE question per turn. Wait for the user's answer before proceeding.\n"
    "- Base every question directly on content in the provided source excerpts.\n"
    "- After the user answers, give brief feedback: correct, partially correct, or incorrect. "
    "Explain using the source and cite it in brackets. Then immediately ask the next question.\n"
    "- Vary question types across the session: definition, application, true/false, "
    "fill-in-the-blank, compare and contrast.\n"
    "- Progress from foundational to more nuanced concepts as the session continues.\n"
    "- After each answer, report the running score (e.g. '2 out of 3 correct so far').\n"
    "- If the user says they don't know or want to skip, give the answer from the source and move on."
)

_CHAT_MODE_OPENINGS = {
    "Interview": (
        "I'd like to explore your family system with you through the lens of Bowen theory. "
        "I'll ask one question at a time — there are no right or wrong answers, "
        "just an opportunity to look at patterns.\n\n"
        "**Where would you like to start?**"
    ),
    "Coach": (
        "I'm here to help you think about your own part in your relationship system — "
        "not to tell you what to do, but to ask questions that help you see your own functioning "
        "more clearly.\n\n"
        "**What situation or pattern in your system is on your mind?**"
    ),
    "Quiz": (
        "Let's test your knowledge of Bowen Family Systems Theory. "
        "I'll ask one question at a time drawn from the source material, "
        "give you feedback after each answer, and track your score.\n\n"
        "Type **start** (or anything) and I'll ask your first question."
    ),
}

CLAUDE_MODELS = [
    "claude-opus-4-7", "claude-sonnet-4-6", "claude-haiku-4-5",
    "claude-opus-4-6", "claude-opus-4-5", "claude-sonnet-4-5",
]
OPENAI_MODELS    = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "o1", "o1-mini"]
DEEPSEEK_MODELS  = ["deepseek-v4-flash", "deepseek-v4-pro"]
DEEPSEEK_BASE_URL = "https://api.deepseek.com/anthropic"


# ══════════════════════════════════════════════════════════════════════════════
# IndexManager  (identical to bowen_rag_gui.py — no tkinter dependency)
# ══════════════════════════════════════════════════════════════════════════════

class IndexManager:
    def __init__(self):
        self.chunks: list = []
        self.matrix       = None
        self.vectorizer   = None
        self.bm25         = None
        self.embed_matrix = None
        self.embed_model  = None
        self.loaded       = False

    def load(self, refs_dir: Path = REFS_DIR) -> dict:
        meta_path  = refs_dir / "chunk_metadata.json"
        matrix_npz = refs_dir / "tfidf_matrix.npz"
        matrix_npy = refs_dir / "tfidf_matrix.npy"

        if not meta_path.exists():
            raise FileNotFoundError(f"Index not found at {refs_dir}.")

        with open(meta_path) as f:
            self.chunks = json.load(f)

        if matrix_npz.exists():
            self.matrix = sp_sparse.load_npz(str(matrix_npz)).toarray()
        else:
            self.matrix = np.load(str(matrix_npy))

        texts = [c["text"] for c in self.chunks]
        self.vectorizer = TfidfVectorizer(
            max_features=8000, stop_words="english",
            lowercase=True, ngram_range=(1, 2),
            min_df=2, sublinear_tf=True
        )
        self.vectorizer.fit(texts)
        self.loaded = True

        self._doc_chunk_ids: dict = {}
        for i, c in enumerate(self.chunks):
            self._doc_chunk_ids.setdefault(c["doc_name"], []).append(i)

        docs = len(set(c["doc_name"] for c in self.chunks))

        embed_npy = refs_dir / "embed_matrix.npy"
        if EMBEDDING_AVAILABLE and embed_npy.exists():
            self.embed_matrix = np.load(str(embed_npy))

        if BM25_AVAILABLE:
            tokenized = [self._tokenize(c["text"]) for c in self.chunks]
            self.bm25 = BM25Okapi(tokenized)

        return {"chunks": len(self.chunks), "documents": docs,
                "embeddings": self.embed_matrix is not None}

    def get_context_window(self, chunk_id: int, window: int = 2) -> list:
        doc_name = self.chunks[chunk_id]["doc_name"]
        doc_ids  = self._doc_chunk_ids.get(doc_name, [])
        try:
            pos = doc_ids.index(chunk_id)
        except ValueError:
            return [self.chunks[chunk_id]["text"]]
        start = max(0, pos - window)
        end   = min(len(doc_ids), pos + window + 1)
        return [self.chunks[doc_ids[j]]["text"] for j in range(start, end)]

    def semantic_search(self, query: str, top_k: int, use_boost: bool = True) -> list:
        if not self.loaded:
            return []
        qvec     = self.vectorizer.transform([query])
        raw      = cosine_similarity(qvec, self.matrix)[0]
        boost_fn = authority_boost if use_boost else (lambda _: 1.0)
        boosted  = np.array([raw[i] * boost_fn(self.chunks[i]["doc_name"])
                             for i in range(len(self.chunks))])
        idx = boosted.argsort()[::-1][:top_k]
        return [
            {**self.chunks[i],
             "score":       float(boosted[i]),
             "score_label": f"{boosted[i]*100:.0f}% ★"
                            if use_boost and authority_boost(self.chunks[i]["doc_name"]) > 1.0
                            else f"{boosted[i]*100:.0f}%",
             "mode": "semantic"}
            for i in idx if boosted[i] > 0
        ]

    @staticmethod
    def _tokenize(text: str) -> list:
        return [w for w in re.split(r'\W+', text.lower())
                if len(w) > 2 and w not in IndexManager._STOP]

    _STOP = frozenset({
        "the","and","for","are","but","not","you","all","can","had","her","was",
        "one","our","out","day","get","has","him","his","how","man","new","now",
        "old","see","two","way","who","boy","did","its","let","put","say","she",
        "too","use","what","with","this","that","have","from","they","will","been",
        "more","when","than","them","were","said","each","which","about","there",
        "their","would","make","like","into","time","look","just","come","could",
        "also","some","then","these","many","well","only","over","such","after",
        "most","very","even","back","any","good","know","same","tell","does",
        "bowen","kerr","theory","family","therapy","systems","system",
        "murray","michael","dr","said","think","know","people","things",
    })

    @staticmethod
    def _stems(word: str) -> list:
        variants = [word]
        if word.endswith("ing") and len(word) > 5:
            variants.append(word[:-3])
        if word.endswith("ed") and len(word) > 4:
            variants.append(word[:-2])
            variants.append(word[:-1])
        if word.endswith("ies") and len(word) > 4:
            variants.append(word[:-3] + "y")
        if word.endswith("es") and len(word) > 4:
            variants.append(word[:-2])
        if word.endswith("s") and len(word) > 3:
            variants.append(word[:-1])
        return variants

    def keyword_search(self, query: str, top_k: int, use_boost: bool = True) -> list:
        if not self.loaded:
            return []

        # Extract quoted phrases ("exact phrase") and remaining individual words
        phrases   = [p.lower() for p in re.findall(r'"([^"]+)"', query)]
        remainder = re.sub(r'"[^"]+"', '', query)
        raw_terms = [t.lower() for t in remainder.split()
                     if len(t) > 2 and t.lower() not in self._STOP]

        if not phrases and not raw_terms:
            return []

        term_sets = [set(self._stems(t)) for t in raw_terms]
        boost_fn  = authority_boost if use_boost else (lambda _: 1.0)
        doc_best: dict = {}

        for c in self.chunks:
            tl = c["text"].lower()

            # Quoted phrases must all be present — skip chunk if any are absent
            if any(tl.count(p) == 0 for p in phrases):
                continue
            phrase_hits = sum(tl.count(p) * 3 for p in phrases)  # weight phrases 3×

            word_hits = sum(max(tl.count(v) for v in variants) for variants in term_sets)
            hits = phrase_hits + word_hits
            if hits == 0:
                continue

            dn    = c["doc_name"]
            boost_val = authority_boost(dn) if use_boost else 1.0
            score = hits * boost_val
            tag   = " phrase" if phrases else ""
            if boost_val > 1.0:
                label = f"{hits} hits{tag} ×{boost_val:g} ★"
            else:
                label = f"{hits} hits{tag}"
            if dn not in doc_best or score > doc_best[dn]["score"]:
                doc_best[dn] = {**c, "score": score, "score_label": label, "mode": "keyword"}

        out = sorted(doc_best.values(), key=lambda x: x["score"], reverse=True)
        return out[:top_k]

    def combined_search(self, query: str, top_k: int, use_boost: bool = True) -> list:
        sem = {r["id"]: r for r in self.semantic_search(query, top_k, use_boost=use_boost)}
        kw  = {r["id"]: r for r in self.keyword_search(query, top_k, use_boost=use_boost)}
        merged = {**kw, **sem}
        return sorted(merged.values(), key=lambda x: x["score"], reverse=True)[:top_k]

    def bm25_search(self, query: str, top_k: int, use_boost: bool = True) -> list:
        if not self.loaded or self.bm25 is None:
            return []
        tokens = self._tokenize(query)
        if not tokens:
            return []
        boost_fn = authority_boost if use_boost else (lambda _: 1.0)
        raw      = self.bm25.get_scores(tokens)
        boosted  = np.array([raw[i] * boost_fn(self.chunks[i]["doc_name"])
                             for i in range(len(self.chunks))])
        idx = boosted.argsort()[::-1][:top_k]
        return [
            {**self.chunks[i],
             "score":       float(boosted[i]),
             "score_label": f"{boosted[i]:.2f} ★"
                            if use_boost and authority_boost(self.chunks[i]["doc_name"]) > 1.0
                            else f"{boosted[i]:.2f}",
             "mode": "bm25"}
            for i in idx if boosted[i] > 0
        ]

    def hybrid_search(self, query: str, top_k: int, use_boost: bool = True) -> list:
        if self.embed_matrix is None:
            raise RuntimeError(
                "Embedding index not built. Run 'Build Embeddings' locally and commit embed_matrix.npy.")
        if len(self.embed_matrix) != len(self.chunks):
            raise RuntimeError(
                f"Embedding index is stale ({len(self.embed_matrix)} vs {len(self.chunks)} chunks).")
        pool         = min(top_k * 4, len(self.chunks))
        bm25_results  = self.bm25_search(query, pool, use_boost=use_boost)
        embed_results = self.embedding_search(query, pool, use_boost=use_boost)
        bm25_rank  = {r["id"]: i for i, r in enumerate(bm25_results)}
        embed_rank = {r["id"]: i for i, r in enumerate(embed_results)}
        K      = 60
        all_ids = set(bm25_rank) | set(embed_rank)
        rrf = {cid: (1.0 / (K + bm25_rank[cid])  if cid in bm25_rank  else 0.0)
                   + (1.0 / (K + embed_rank[cid]) if cid in embed_rank else 0.0)
               for cid in all_ids}
        top_ids      = sorted(rrf, key=lambda x: rrf[x], reverse=True)[:top_k]
        chunk_lookup = {r["id"]: r for r in embed_results + bm25_results}
        max_rrf      = 2.0 / K
        results = []
        for cid in top_ids:
            base = chunk_lookup[cid].copy()
            pct  = rrf[cid] / max_rrf * 100
            boosted_flag = use_boost and authority_boost(base["doc_name"]) > 1.0
            base["score"]       = rrf[cid]
            base["score_label"] = f"{pct:.0f}% ⬡" + (" ★" if boosted_flag else "")
            base["mode"]        = "hybrid"
            results.append(base)
        return results

    def embedding_search(self, query: str, top_k: int, use_boost: bool = True) -> list:
        if not self.loaded or self.embed_matrix is None:
            raise RuntimeError(
                "Embedding index not available. Build it locally and commit embed_matrix.npy.")
        if self.embed_model is None:
            self.embed_model = SentenceTransformer("all-MiniLM-L6-v2")
        if len(self.embed_matrix) != len(self.chunks):
            raise RuntimeError("Embedding index is stale — rebuild it.")
        qvec    = self.embed_model.encode([query])
        raw     = cosine_similarity(qvec, self.embed_matrix)[0]
        boost_fn = authority_boost if use_boost else (lambda _: 1.0)
        boosted = np.array([raw[i] * boost_fn(self.chunks[i]["doc_name"])
                            for i in range(len(self.chunks))])
        idx = boosted.argsort()[::-1][:top_k]
        return [
            {**self.chunks[i],
             "score":       float(boosted[i]),
             "score_label": f"{boosted[i]*100:.0f}% ✦"
                            if use_boost and authority_boost(self.chunks[i]["doc_name"]) > 1.0
                            else f"{boosted[i]*100:.0f}%",
             "mode": "embedding"}
            for i in idx if boosted[i] > 0
        ]

    def top_docs_search(self, query: str, top_chunks: int = 300,
                        top_docs: int = 30, use_boost: bool = True) -> list:
        if not self.loaded:
            return []
        qvec     = self.vectorizer.transform([query])
        raw      = cosine_similarity(qvec, self.matrix)[0]
        boost_fn = authority_boost if use_boost else (lambda _: 1.0)
        doc_chunks: dict = {}
        for i, score in enumerate(raw):
            if score <= 0:
                continue
            dn = self.chunks[i]["doc_name"]
            doc_chunks.setdefault(dn, []).append((score, i))
        doc_scores: dict = {}
        for dn, pairs in doc_chunks.items():
            pairs.sort(reverse=True)
            top3_sum  = sum(s for s, _ in pairs[:3])
            agg_score = top3_sum * boost_fn(dn)
            best_idx  = pairs[0][1]
            label     = f"{agg_score*100:.0f}%"
            if use_boost and authority_boost(dn) > 1.0:
                label += " ★"
            doc_scores[dn] = {
                **self.chunks[best_idx],
                "score":       agg_score,
                "score_label": label,
                "mode":        "semantic",
            }
        out = sorted(doc_scores.values(), key=lambda x: x["score"], reverse=True)
        return out[:top_docs]

    def list_documents(self) -> list:
        seen, docs = set(), []
        for c in self.chunks:
            if c["doc_name"] not in seen:
                seen.add(c["doc_name"])
                docs.append(c["doc_name"])
        return sorted(docs)


# ══════════════════════════════════════════════════════════════════════════════
# LLM streaming
# ══════════════════════════════════════════════════════════════════════════════

def _llm_stream(messages: list, system: str):
    """Generator that yields tokens from the configured LLM provider."""
    ss       = st.session_state
    provider = ss.get("provider", "claude")

    if provider == "claude":
        key = ss.get("claude_key", "")
        if not key:
            raise RuntimeError("Claude API key not set — go to Settings.")
        import anthropic
        client = anthropic.Anthropic(api_key=key)
        with client.messages.stream(
            model=ss.get("claude_model", "claude-sonnet-4-6"),
            max_tokens=16000, system=system, messages=messages
        ) as s:
            for token in s.text_stream:
                yield token

    elif provider == "openai":
        key = ss.get("openai_key", "")
        if not key:
            raise RuntimeError("OpenAI API key not set — go to Settings.")
        import openai
        client = openai.OpenAI(api_key=key)
        full = [{"role": "system", "content": system}] + messages
        with client.chat.completions.create(
            model=ss.get("openai_model", "gpt-4o"),
            max_tokens=16000, messages=full, stream=True
        ) as s:
            for chunk in s:
                t = chunk.choices[0].delta.content or ""
                if t:
                    yield t

    elif provider == "deepseek":
        key = ss.get("deepseek_key", "")
        if not key:
            raise RuntimeError("DeepSeek API key not set — go to Settings.")
        import anthropic
        client = anthropic.Anthropic(api_key=key, base_url=DEEPSEEK_BASE_URL)
        with client.messages.stream(
            model=ss.get("deepseek_model", "deepseek-v4-flash"),
            max_tokens=8000, system=system, messages=messages
        ) as s:
            for token in s.text_stream:
                yield token

    else:  # ollama
        import requests
        url   = ss.get("ollama_url", "http://localhost:11434")
        model = ss.get("ollama_model", "qwen2.5:7b")
        full  = [{"role": "system", "content": system}] + messages
        try:
            r = requests.post(f"{url.rstrip('/')}/api/chat",
                              json={"model": model, "messages": full, "stream": True},
                              stream=True, timeout=300)
            r.raise_for_status()
        except requests.exceptions.ConnectionError:
            raise RuntimeError(
                f"Cannot connect to Ollama at {url}. "
                "Ollama must be running locally to use this provider. "
                "Go to Settings and switch to DeepSeek, Claude, or OpenAI."
            )
        for line in r.iter_lines():
            if line:
                d = json.loads(line)
                t = d.get("message", {}).get("content", "")
                if t:
                    yield t
                if d.get("done"):
                    break


# ══════════════════════════════════════════════════════════════════════════════
# App bootstrap
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Loading index…")
def _get_index() -> IndexManager:
    idx = IndexManager()
    idx.load(REFS_DIR)
    return idx


def _init_session():
    defaults = {
        "search_results":  [],
        "staged_chunks":   [],
        "chat_history":    [],
        "last_rpt_context":  "",
        "last_rpt_appendix": "",
        "last_report":       "",
        "provider":        os.environ.get("LLM_PROVIDER", "deepseek"),
        "claude_key":      os.environ.get("ANTHROPIC_API_KEY", ""),
        "claude_model":    os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
        "openai_key":      os.environ.get("OPENAI_API_KEY", ""),
        "openai_model":    os.environ.get("OPENAI_MODEL", "gpt-4o"),
        "ollama_url":      os.environ.get("OLLAMA_URL", "http://localhost:11434"),
        "ollama_model":    os.environ.get("OLLAMA_MODEL", "qwen2.5:7b"),
        "deepseek_key":    os.environ.get("DEEPSEEK_API_KEY", ""),
        "deepseek_model":  os.environ.get("DEEPSEEK_MODEL", "deepseek-v4-flash"),
        "system_prompt":           SYSTEM_PROMPT,
        "default_search_mode":     "hybrid",
        "chat_interaction_mode":   "Standard",
        "citation_style": citations.normalize_style(
            os.environ.get("CITATION_STYLE", citations.DEFAULT_STYLE)),
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # Env vars always override stale session state for provider and keys.
    # This prevents a browser session that had "ollama" selected from
    # persisting that choice after a redeploy changes LLM_PROVIDER.
    _valid_providers = {"claude", "openai", "deepseek", "ollama"}
    raw_provider = os.environ.get("LLM_PROVIDER", "")
    if raw_provider:
        # Guard against LLM_PROVIDER being set to a model name (e.g. "deepseek-v4-flash")
        st.session_state["provider"] = raw_provider if raw_provider in _valid_providers else "deepseek"
    for env_var, ss_key in (
        ("ANTHROPIC_API_KEY", "claude_key"),
        ("OPENAI_API_KEY",    "openai_key"),
        ("DEEPSEEK_API_KEY",  "deepseek_key"),
    ):
        if env_var in os.environ:
            st.session_state[ss_key] = os.environ[env_var]


def _check_auth():
    required = os.environ.get("APP_PASSWORD", "")
    if not required:
        return
    if st.session_state.get("authenticated"):
        return
    _, col, _ = st.columns([1, 2, 1])
    with col:
        st.title("Bowen Theory RAG")
        st.caption("Enter the access password to continue.")
        pwd = st.text_input("Password", type="password")
        if st.button("Enter", type="primary", use_container_width=True):
            if pwd == required:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("Incorrect password.")
    st.stop()


def _gather_chunks(idx: IndexManager, query: str) -> list:
    """If staged chunks exist, use only those. Otherwise, retrieve fresh."""
    ss     = st.session_state
    staged = ss.get("staged_chunks", [])

    # If user staged chunks, use ONLY those (don't retrieve fresh)
    if staged:
        return staged

    # Otherwise, retrieve fresh chunks based on query
    mode      = ss.get("rpt_mode", "top-docs")
    k         = ss.get("rpt_k", 30)
    use_boost = ss.get("rpt_boost", True)

    if "top-docs" in mode:
        return idx.top_docs_search(query, top_chunks=300, top_docs=k, use_boost=use_boost)
    elif mode == "semantic":
        return idx.semantic_search(query, k, use_boost=use_boost)
    elif mode == "keyword":
        return idx.keyword_search(query, k, use_boost=use_boost)
    elif mode == "embedding":
        return idx.embedding_search(query, k, use_boost=use_boost)
    elif mode == "hybrid":
        return idx.hybrid_search(query, k, use_boost=use_boost)
    else:
        return idx.combined_search(query, k, use_boost=use_boost)


def _format_chunk_text(text: str) -> str:
    """Return chunk text cleaned up for readable markdown display."""
    # Strip [Section Title]\n\n prefix added by the indexer
    text = re.sub(r'^\[[^\]]+\]\s*\n\n', '', text, count=1)
    # Split on paragraph breaks (2+ newlines)
    paragraphs = re.split(r'\n{2,}', text)
    cleaned = []
    for para in paragraphs:
        # Collapse soft-wrap newlines within a paragraph (common in PDF text)
        para = re.sub(r'[ \t]*\n[ \t]*', ' ', para)
        # Collapse multiple spaces
        para = re.sub(r' {2,}', ' ', para).strip()
        if para:
            cleaned.append(para)
    return '\n\n'.join(cleaned)


# ─── Export converters ──────────────────────────────────────────────────────
def md_to_docx_bytes(md_text: str, title: str = "Bowen Report") -> bytes:
    """Convert markdown text to a .docx file, returned as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from io import BytesIO

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for raw_line in md_text.split('\n'):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        # Bullet lists
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
            text = line.lstrip()[2:].strip()
            doc.add_paragraph(_clean_md_inline(text), style='List Bullet')
        # Numbered lists (1. 2. etc.)
        elif re.match(r'^\d+\.\s', line.lstrip()):
            text = re.sub(r'^\d+\.\s', '', line.lstrip())
            doc.add_paragraph(_clean_md_inline(text), style='List Number')
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('_' * 40)
        else:
            p = doc.add_paragraph()
            _add_runs_with_formatting(p, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _clean_md_inline(text: str) -> str:
    """Strip basic markdown inline formatting markers."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def _add_runs_with_formatting(paragraph, text: str) -> None:
    """Add text to a docx paragraph, handling **bold** and *italic* inline."""
    # Pattern to split on **bold** and *italic*
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def md_to_pdf_bytes(md_text: str, title: str = "Bowen Report") -> bytes:
    """Convert markdown text to a PDF file, returned as bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.lib.enums import TA_LEFT
    from io import BytesIO

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                fontSize=11, leading=15, spaceAfter=6)
    h1_style = ParagraphStyle('H1', parent=styles['Heading1'],
                              fontSize=18, leading=22, spaceBefore=12, spaceAfter=8)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'],
                              fontSize=15, leading=19, spaceBefore=10, spaceAfter=6)
    h3_style = ParagraphStyle('H3', parent=styles['Heading3'],
                              fontSize=13, leading=17, spaceBefore=8, spaceAfter=4)

    story = []
    for raw_line in md_text.split('\n'):
        line = raw_line.rstrip()
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
        # Escape HTML special chars for reportlab
        def fmt(t):
            t = (t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
            # Bold
            t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
            # Italic
            t = re.sub(r'\*([^*]+?)\*', r'<i>\1</i>', t)
            return t

        if line.startswith('### '):
            story.append(Paragraph(fmt(line[4:].strip()), h3_style))
        elif line.startswith('## '):
            story.append(Paragraph(fmt(line[3:].strip()), h2_style))
        elif line.startswith('# '):
            story.append(Paragraph(fmt(line[2:].strip()), h1_style))
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
            text = line.lstrip()[2:].strip()
            story.append(Paragraph('• ' + fmt(text), body_style))
        elif re.match(r'^\d+\.\s', line.lstrip()):
            story.append(Paragraph(fmt(line.lstrip()), body_style))
        elif line.strip() == '---':
            story.append(Spacer(1, 6))
            story.append(Paragraph('_' * 60, body_style))
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(fmt(line), body_style))

    doc.build(story)
    return buf.getvalue()


@st.dialog("Section text", width="large")
def _show_section_dialog() -> None:
    result = st.session_state.get("_dialog_result", {})
    doc    = result.get("doc_name", "")
    title  = result.get("section_title", "")
    page   = result.get("page")
    chunk_pos   = result.get("chunk_pos")
    doc_total   = result.get("doc_chunk_count", 1)

    parts = [doc]
    if title:
        parts.append(title)
    if page:
        parts.append(f"p. {page}")
    elif chunk_pos and doc_total > 1:
        parts.append(f"~{round(chunk_pos / doc_total * 100)}% through document")
    st.caption(" · ".join(parts))
    st.divider()
    st.markdown(_format_chunk_text(result.get("text", "")))


def _score_color(result: dict) -> str:
    score = result.get("score", 0)
    if result.get("mode") == "keyword":
        return "#7c3aed"
    if isinstance(score, float):
        if score > 0.4:
            return "#16a34a"
        if score > 0.15:
            return "#ca8a04"
    return "#6b7280"


def _result_card(result: dict, checkbox_key: str):
    """Render a single search result card with checkbox."""
    author  = doc_author(result["doc_name"])
    color   = _score_color(result)
    excerpt = re.sub(r'\s+', ' ', result["text"][:220]).strip()
    page    = result.get("page")

    col_cb, col_body, col_view = st.columns([0.04, 0.88, 0.08])
    with col_cb:
        checked = st.checkbox("Select", key=checkbox_key, label_visibility="collapsed")
    with col_body:
        badges = (
            f'<span style="background:{color};color:white;padding:2px 8px;'
            f'border-radius:4px;font-size:11px;margin-right:4px">{result["score_label"]}</span>'
        )
        if author != "Unknown":
            badges += (
                f'<span style="background:#7c3aed;color:white;padding:2px 8px;'
                f'border-radius:4px;font-size:11px;margin-right:4px">{author}</span>'
            )
        if page:
            badges += (
                f'<span style="background:#475569;color:white;padding:2px 8px;'
                f'border-radius:4px;font-size:11px;margin-right:4px">p.{page}</span>'
            )
        else:
            chunk_pos = result.get("chunk_pos")
            doc_total = result.get("doc_chunk_count", 1)
            if chunk_pos and doc_total > 1:
                pct = round(chunk_pos / doc_total * 100)
                badges += (
                    f'<span style="background:#64748b;color:white;padding:2px 8px;'
                    f'border-radius:4px;font-size:11px;margin-right:4px">~{pct}%</span>'
                )
        st.markdown(
            f'{badges}<strong style="font-size:13px">{result["doc_name"]}</strong>',
            unsafe_allow_html=True
        )
        st.caption(excerpt + "…")
    with col_view:
        if st.button("↗", key=f"view_{checkbox_key}", use_container_width=True):
            st.session_state["_dialog_result"] = result
            _show_section_dialog()
    return checked


# ══════════════════════════════════════════════════════════════════════════════
# Pages
# ══════════════════════════════════════════════════════════════════════════════

def page_search(idx: IndexManager):
    st.header("Search")

    # 40% controls panel | 60% results — wide enough for help= icons to render
    ctrl, results_col = st.columns([2, 3])

    with ctrl:
        query = st.text_area(
            "Query", height=100, placeholder="Enter your search query…",
            help="Enter keywords or a natural-language question about Bowen theory.",
        )

        mode_options = [
            ("Top Docs (recommended)", "top-docs"),
            ("Semantic (TF-IDF)",      "semantic"),
            ("Keyword",                "keyword"),
            ("Both",                   "both"),
        ]
        if EMBEDDING_AVAILABLE and idx.embed_matrix is not None:
            mode_options.append(("Embedding", "embedding"))
        if EMBEDDING_AVAILABLE and BM25_AVAILABLE and idx.embed_matrix is not None:
            mode_options.append(("Hybrid (BM25 + Embedding)", "hybrid"))

        mode_labels = [m[0] for m in mode_options]
        mode_values = [m[1] for m in mode_options]
        default_mode = st.session_state.get("default_search_mode", "hybrid")
        default_idx  = mode_values.index(default_mode) if default_mode in mode_values else 0

        mode_idx = st.selectbox(
            "Mode", range(len(mode_labels)),
            format_func=lambda i: mode_labels[i],
            index=default_idx,
            help=(
                "**Top Docs** — aggregates per document; best for most queries.\n\n"
                "**Semantic** — TF-IDF cosine similarity; fast, exact-vocabulary.\n\n"
                "**Keyword** — exact word matching with stemming; good for names.\n\n"
                "**Both** — merges semantic + keyword.\n\n"
                "**Embedding** — conceptual similarity regardless of exact wording.\n\n"
                "**Hybrid** — BM25 + Embedding via Reciprocal Rank Fusion; usually best overall."
            ),
        )
        mode = mode_values[mode_idx]

        top_k = st.number_input(
            "Results", min_value=1, max_value=200, value=15,
            help="Maximum number of results to return.",
        )
        use_boost = st.checkbox(
            "Authority boost", value=True,
            help=(
                "This prioritizes Bowen, Kerr, Papero as a source, but does not eliminate other sources. "
                "Primary Bowen/Kerr sources (3×), FSJ articles (1.3×), and other named theorists (1.15×)."
            ),
        )

        authors       = ["All authors"] + all_known_authors()
        author_filter = st.selectbox(
            "Author filter", authors,
            help="Narrow results to a specific author.",
        )

        search_clicked = st.button("Search", type="primary", use_container_width=True)

        st.divider()
        staged = st.session_state.get("staged_chunks", [])
        if staged:
            st.success(f"{len(staged)} chunks staged for Report")
            if st.button("Clear staged", use_container_width=True):
                st.session_state.staged_chunks = []
                st.rerun()

    if search_clicked and query.strip():
        st.session_state.last_search_query = query.strip()
        with st.spinner("Searching…"):
            try:
                if mode == "top-docs":
                    results = idx.top_docs_search(query, top_chunks=300,
                                                  top_docs=top_k, use_boost=use_boost)
                elif mode == "semantic":
                    results = idx.semantic_search(query, top_k, use_boost=use_boost)
                elif mode == "keyword":
                    results = idx.keyword_search(query, top_k, use_boost=use_boost)
                elif mode == "embedding":
                    results = idx.embedding_search(query, top_k, use_boost=use_boost)
                elif mode == "hybrid":
                    results = idx.hybrid_search(query, top_k, use_boost=use_boost)
                else:
                    results = idx.combined_search(query, top_k, use_boost=use_boost)
            except RuntimeError as e:
                st.error(str(e))
                results = []

        if author_filter != "All authors":
            results = [r for r in results if doc_author(r["doc_name"]) == author_filter]

        for key in list(st.session_state.keys()):
            if key.startswith("sel_"):
                del st.session_state[key]

        st.session_state.search_results = results
        # Sync boost setting to Report tab
        st.session_state.rpt_boost = use_boost

    with results_col:
        results = st.session_state.get("search_results", [])
        if not results:
            st.info("Run a search to see results.")
        else:
            sel_col1, sel_col2, sel_col3, sel_col4 = st.columns([1, 1, 1.5, 0.5])
            with sel_col1:
                if st.button("Select All"):
                    for i, r in enumerate(results):
                        st.session_state[f"sel_{r.get('id', i)}"] = True
                    st.rerun()
            with sel_col2:
                if st.button("Clear selection"):
                    for i, r in enumerate(results):
                        st.session_state[f"sel_{r.get('id', i)}"] = False
                    st.rerun()
            with sel_col3:
                if st.button(
                    "Stage selected for Report", type="primary",
                    help="Save checked results to include in a Report. Optional — you can also go to Report and do a fresh retrieval.",
                ):
                    selected = [results[i] for i, r in enumerate(results)
                                if st.session_state.get(f"sel_{r.get('id', i)}", False)]
                    if selected:
                        st.session_state.staged_chunks = selected
                        st.success(f"{len(selected)} chunks staged.")
                    else:
                        st.warning("Select at least one result first.")
            with sel_col4:
                if st.button("?", key="help_stage_workflow", use_container_width=True):
                    st.session_state["show_stage_help"] = not st.session_state.get("show_stage_help", False)
                    st.rerun()

            if st.session_state.get("show_stage_help", False):
                st.info(
                    "**Report Workflow Options:**\n\n"
                    "**Option 1: Fresh Retrieval (recommended for most queries)**\n"
                    "1. Run a search to find relevant sources\n"
                    "2. Go to the Report tab\n"
                    "3. Enter your topic and click \"Generate Report\"\n"
                    "4. Report retrieves fresh chunks for that query\n\n"
                    "**Option 2: Use Selected Results**\n"
                    "1. Run a search\n"
                    "2. Use Select/Clear to filter results\n"
                    "3. Click \"Stage selected for Report\" to save them\n"
                    "4. Go to Report tab and generate\n"
                    "5. Report combines your staged chunks + fresh retrieval\n\n"
                    "**When to stage:**\n"
                    "✓ You found specific sources you want to highlight\n"
                    "✓ You want to pre-filter low-quality results\n"
                    "✗ You just want fresh results—no need to stage"
                )

            _rc, _hc = st.columns([9, 1])
            with _rc:
                st.caption(f"{len(results)} results")
            with _hc:
                if st.button("?", key="help_score_badges", use_container_width=True):
                    st.session_state["show_help_badges"] = not st.session_state.get("show_help_badges", False)
                    st.rerun()
            if st.session_state.get("show_help_badges", False):
                st.info(
                    "**Score badges**\n\n"
                    "| Badge | Mode | Meaning |\n"
                    "|---|---|---|\n"
                    "| `45%` | Top Docs | Aggregate similarity across best chunks for that document |\n"
                    "| `45% ★` | Semantic (TF-IDF) | Cosine similarity; ★ = authority-boosted source |\n"
                    "| `45% ⬡` | Hybrid | Reciprocal Rank Fusion (RRF) score; ⬡ marks Hybrid mode |\n"
                    "| `45% ⬡ ★` | Hybrid + boosted | RRF score with authority boost applied |\n"
                    "| `45% ✦` | Embedding | Cosine similarity on sentence-transformer vectors |\n"
                    "| `6 hits ×3 ★` | Keyword | Raw hit count × authority multiplier |\n"
                    "| `3.42 ★` | BM25 | BM25 relevance score |\n\n"
                    "**Authority boost (★):** Primary Bowen/Kerr sources ×3, "
                    "Family Systems Journal ×1.3, named theorists ×1.15. "
                    "Toggle off with the **Authority boost** checkbox to rank by raw score only."
                )
            st.divider()

            for i, result in enumerate(results):
                cb_key = f"sel_{result.get('id', i)}"
                _result_card(result, cb_key)


def page_chat(idx: IndexManager):
    st.header("Chat")

    # ── Interaction mode selector ─────────────────────────────────────────────
    _interaction_modes = ["Standard", "Interview", "Coach", "Quiz"]
    _prev_mode = st.session_state.get("chat_interaction_mode", "Standard")
    _mode_descriptions = {
        "Standard":  "Ask questions; the AI answers from the source corpus.",
        "Interview": "The AI interviews you about your family system using Bowen concepts.",
        "Coach":     "The AI coaches you in applying Bowen theory to your own functioning.",
        "Quiz":      "The AI quizzes you on Bowen theory, one question at a time with scoring.",
    }
    im_col, imdesc_col = st.columns([2, 5])
    with im_col:
        _selected_mode = st.selectbox(
            "Mode",
            _interaction_modes,
            index=_interaction_modes.index(_prev_mode),
            key="chat_interaction_mode_sel",
        )
    with imdesc_col:
        st.write("")
        st.caption(_mode_descriptions[_selected_mode])

    if _selected_mode != _prev_mode:
        st.session_state["chat_interaction_mode"] = _selected_mode
        opening = _CHAT_MODE_OPENINGS.get(_selected_mode)
        st.session_state.chat_history = (
            [{"role": "assistant", "content": opening, "sources": []}]
            if opening else []
        )
        st.rerun()

    _chat_mode = st.session_state.get("chat_interaction_mode", "Standard")

    st.divider()

    # Single compact row — retrieval options
    cc1, cc2, cc3, cc4, cc5, cc6 = st.columns([2, 1, 1, 2, 1, 1])

    chat_mode_opts = ["top-docs", "semantic", "keyword", "both"]
    if EMBEDDING_AVAILABLE and idx.embed_matrix is not None:
        chat_mode_opts.insert(0, "embedding")
        if BM25_AVAILABLE:
            chat_mode_opts.insert(1, "hybrid")
    default_mode     = st.session_state.get("default_search_mode", "hybrid")
    chat_default_idx = chat_mode_opts.index(default_mode) if default_mode in chat_mode_opts else 0

    with cc1:
        chat_mode = st.selectbox("Mode", chat_mode_opts, index=chat_default_idx,
                                 key="chat_mode_sel")
    with cc2:
        chat_k = st.number_input("Chunks", min_value=3, max_value=100, value=12,
                                 key="chat_k_inp")
    with cc3:
        chat_boost = st.checkbox(
            "Boost", value=True, key="chat_boost_cb",
            help=(
                "This prioritizes Bowen, Kerr, Papero as a source, but does not eliminate other sources. "
                "Primary Bowen/Kerr sources 3×, FSJ articles 1.3×, other named theorists 1.15×."
            ),
        )
    with cc4:
        chat_authors = ["All authors"] + all_known_authors()
        chat_author  = st.selectbox("Author", chat_authors, key="chat_author_sel")
    with cc5:
        st.write("")
        if st.button("Clear", use_container_width=True):
            st.session_state.chat_history = []
            st.rerun()
    with cc6:
        st.write("")
        if st.button("?", key="chat_help_btn", use_container_width=True):
            st.session_state["show_chat_help"] = not st.session_state.get("show_chat_help", False)

    if st.session_state.get("show_chat_help", False):
        st.info(
            "**Mode** — how source chunks are retrieved "
            "(Hybrid = BM25 + Embedding, best overall).\n\n"
            "**Chunks** — number of source passages retrieved per question.\n\n"
            "**Boost** — apply authority weighting (primary sources rank 3× higher).\n\n"
            "**Author** — restrict retrieval to a specific author's works."
        )

    # Display history
    for msg_idx, msg in enumerate(st.session_state.chat_history):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant" and msg.get("sources"):
                with st.expander(f"Sources ({len(msg['sources'])} docs)"):
                    for src_idx, src in enumerate(msg["sources"]):
                        sc, vc = st.columns([5, 1])
                        with sc:
                            st.caption(f"**{src['doc']}** — {src['excerpt']}…")
                        with vc:
                            chunk = next((c for c in idx.chunks
                                          if c["doc_name"] == src["doc"]), None)
                            if chunk and st.button("↗",
                                                   key=f"chat_view_{msg_idx}_{src_idx}"):
                                st.session_state["_dialog_result"] = chunk
                                _show_section_dialog()

    # Input
    _placeholders = {
        "Standard":  "Ask about Bowen theory…",
        "Interview": "Respond to the question above…",
        "Coach":     "Share what's on your mind…",
        "Quiz":      "Type your answer…",
    }
    if prompt := st.chat_input(_placeholders.get(_chat_mode, "Ask about Bowen theory…")):
        with st.chat_message("user"):
            st.markdown(prompt)

        # Retrieve chunks
        try:
            if chat_mode == "top-docs":
                chunks = idx.top_docs_search(prompt, top_chunks=300,
                                             top_docs=chat_k, use_boost=chat_boost)
            elif chat_mode == "semantic":
                chunks = idx.semantic_search(prompt, chat_k, use_boost=chat_boost)
            elif chat_mode == "keyword":
                chunks = idx.keyword_search(prompt, chat_k, use_boost=chat_boost)
            elif chat_mode == "embedding":
                chunks = idx.embedding_search(prompt, chat_k, use_boost=chat_boost)
            elif chat_mode == "hybrid":
                chunks = idx.hybrid_search(prompt, chat_k, use_boost=chat_boost)
            else:
                chunks = idx.combined_search(prompt, chat_k, use_boost=chat_boost)
        except RuntimeError as e:
            st.error(str(e))
            return

        if chat_author != "All authors":
            chunks = [c for c in chunks if doc_author(c["doc_name"]) == chat_author]
            if not chunks:
                st.error(f"No chunks found for author: {chat_author}")
                return

        # Build context
        docs: dict = {}
        for c in chunks:
            cid  = c.get("id")
            txts = (idx.get_context_window(cid, window=1)
                    if cid is not None and hasattr(idx, "_doc_chunk_ids")
                    else [c["text"]])
            seen = set(docs.get(c["doc_name"], []))
            for t in txts:
                if t not in seen:
                    docs.setdefault(c["doc_name"], []).append(t)
                    seen.add(t)

        context = "\n\n---\n\n".join(
            f"### [{dn}]\n" + "\n…\n".join(txts) for dn, txts in docs.items()
        )
        current_content = (
            f"[Retrieved {len(chunks)} chunks from {len(docs)} documents]\n\n"
            f"{context}\n\n---\nQuestion: {prompt}"
        )
        messages_to_send = [{"role": m["role"], "content": m["content"]}
                            for m in st.session_state.chat_history] + [
            {"role": "user", "content": current_content}
        ]

        _mode_systems = {
            "Interview": INTERVIEW_SYSTEM_PROMPT,
            "Coach":     COACH_SYSTEM_PROMPT,
            "Quiz":      QUIZ_SYSTEM_PROMPT,
        }
        system = _mode_systems.get(
            st.session_state.get("chat_interaction_mode", "Standard"),
            st.session_state.get("system_prompt", SYSTEM_PROMPT),
        )

        with st.chat_message("assistant"):
            try:
                response = st.write_stream(_llm_stream(messages_to_send, system))
            except Exception as e:
                st.error(f"LLM error: {e}")
                return

            doc_names = sorted(set(c["doc_name"] for c in chunks))
            sources   = [{"doc": d,
                          "excerpt": re.sub(r'\s+', ' ', next(
                              (c["text"][:150] for c in chunks if c["doc_name"] == d), ""))
                          } for d in doc_names]
            # Key index this message will have once appended to history:
            # chat_history currently has N items; user appended first → N,
            # assistant appended second → N+1.
            cur_msg_idx = len(st.session_state.chat_history) + 1
            with st.expander(f"Sources used ({len(doc_names)} docs)"):
                for src_idx, src in enumerate(sources):
                    sc, vc = st.columns([5, 1])
                    with sc:
                        st.caption(f"**{src['doc']}** — {src['excerpt']}…")
                    with vc:
                        chunk = next((c for c in idx.chunks
                                      if c["doc_name"] == src["doc"]), None)
                        if chunk and st.button("↗",
                                               key=f"chat_view_{cur_msg_idx}_{src_idx}"):
                            st.session_state["_dialog_result"] = chunk
                            _show_section_dialog()

        # Bare Q&A stored in history (no chunks)
        st.session_state.chat_history.append({"role": "user",      "content": prompt})
        st.session_state.chat_history.append({"role": "assistant", "content": response,
                                               "sources": sources})


def page_report(idx: IndexManager):
    st.header("Report Generator")

    query = st.text_area(
        "Topic / Question", height=80,
        value=st.session_state.get("last_search_query", ""),
        placeholder='e.g. "What does Bowen theory say about triangles?"',
        help="The question or topic the report will address. Pre-filled from your last search.",
    )

    # Single row: Retrieve | Mode | Target words | Chunks per source
    rc1, rc2, rc3, rc4 = st.columns([1, 2, 1, 1])
    with rc1:
        rpt_k = st.number_input(
            "Retrieve top", min_value=5, max_value=150, value=30,
            help=(
                "How many source chunks to retrieve and pass to the LLM. "
                "More = broader coverage but slower and costlier. "
                "30 is a good starting point; use 50+ for broad topics."
            ),
        )
        st.session_state.rpt_k = rpt_k
    with rc2:
        rpt_mode_opts = ["top-docs (recommended)", "semantic", "keyword", "both"]
        if EMBEDDING_AVAILABLE and idx.embed_matrix is not None:
            rpt_mode_opts.append("embedding")
            if BM25_AVAILABLE:
                rpt_mode_opts.append("hybrid")
        default_mode    = st.session_state.get("default_search_mode", "hybrid")
        rpt_mode_values = [o.split(" ")[0] for o in rpt_mode_opts]
        rpt_default_idx = rpt_mode_values.index(default_mode) if default_mode in rpt_mode_values else 0
        rpt_mode = st.selectbox(
            "Mode", rpt_mode_opts, index=rpt_default_idx,
            help=(
                "Search method used to retrieve source chunks.\n\n"
                "**Hybrid** — best overall coverage.\n\n"
                "**Top Docs** — fast; good for well-indexed topics.\n\n"
                "**Keyword** — good for names and specific terms."
            ),
        )
        st.session_state.rpt_mode = rpt_mode.split(" ")[0]
    with rc3:
        target_words = st.number_input(
            "Target words", min_value=500, max_value=10000, value=2000, step=500,
            help=(
                "Approximate minimum word count for the generated report. "
                "2000 is a solid summary; 4000+ for a deep dive."
            ),
        )
    with rc4:
        cpd = st.number_input(
            "Chunks per source", min_value=1, max_value=20, value=5,
            help=(
                "How many text chunks from each document are included as context "
                "(sliding window around the top chunk). "
                "Higher = more context per document but larger LLM input."
            ),
        )

    cb1, cb2 = st.columns(2)
    with cb1:
        # Use session state as the source of truth for boost setting
        st.session_state.rpt_boost = st.checkbox(
            "Authority boost",
            value=st.session_state.get("rpt_boost", True),
            help=(
                "This prioritizes Bowen, Kerr, Papero as a source, but does not eliminate other sources. "
                "Primary Bowen/Kerr sources 3×, FSJ articles 1.3×, other named theorists 1.15×."
            ),
        )
    with cb2:
        include_appendix = st.checkbox(
            "Include sources as Appendix",
            help="Appends the full formatted text of every cited source after the report body.",
        )

    staged = st.session_state.get("staged_chunks", [])
    if staged:
        st.info(f"{len(staged)} chunks staged from Search will make up this report.")
        if st.button("Clear staged"):
            st.session_state.staged_chunks = []
            st.rerun()

    col_gen, col_clear = st.columns([1, 4])
    with col_gen:
        generate = st.button("Generate Report", type="primary")

    if generate:
        if not query.strip():
            st.warning("Enter a topic or question.")
            return

        with st.spinner("Gathering sources…"):
            try:
                chunks = _gather_chunks(idx, query)
            except RuntimeError as e:
                st.error(str(e))
                return

        if not chunks:
            st.warning("No relevant sources found.")
            return

        # Build context
        docs: dict = {}
        window = max(0, (cpd - 1) // 2)
        for c in chunks:
            cid = c.get("id")
            expanded = (idx.get_context_window(cid, window=window)
                        if cid is not None and hasattr(idx, "_doc_chunk_ids")
                        else [c["text"]])
            existing = set(docs.get(c["doc_name"], []))
            for t in expanded:
                if t not in existing:
                    docs.setdefault(c["doc_name"], []).append(t)
                    existing.add(t)

        ref_map  = {name: i + 1 for i, name in enumerate(sorted(docs))}
        # Page per document, but ONLY when the retrieved chunks for that doc resolve to
        # a SINGLE page — otherwise a page locator would be confidently wrong (a doc's
        # retrieved chunks span several pages). Ambiguous → no page; the model cites
        # without a locator rather than with a made-up one.
        doc_pages: dict = {}
        for c in chunks:
            p = c.get("page")
            if p is not None:
                doc_pages.setdefault(c["doc_name"], set()).add(p)
        doc_page = {dn: next(iter(ps)) for dn, ps in doc_pages.items() if len(ps) == 1}
        # Source list handed to the model — [[N]] so citations can't collide with the
        # single brackets that appear inside quoted source text.
        refs_md  = "\n".join(f"[[{num}]] {name}"
                             for name, num in sorted(ref_map.items(), key=lambda x: x[1]))
        # Deterministic bibliographic record per cited number (sources.yml → filename fallback).
        num_to_record = {ref_map[dn]: citations.record_for_doc(dn, SOURCES, doc_author)
                         for dn in docs}
        context_parts = [
            f"### [[{ref_map[dn]}]] {dn}"
            + (f" (p. {doc_page[dn]})" if dn in doc_page else "") + "\n"
            + "\n…\n".join(txts)
            for dn, txts in docs.items()
        ]
        context = "\n\n---\n\n".join(context_parts)
        st.session_state.last_rpt_context = context

        prompt = f"""Write a comprehensive report on the following topic using ONLY the source excerpts provided below.

**Topic / Question:** {query}

---

## SOURCE EXCERPTS ({len(docs)} documents)

{context}

---

## STRICT INSTRUCTIONS

- **Use only the excerpts above.** Do not add any information from outside these sources.
- **Do not infer, assume, or extrapolate.** If the sources do not explicitly address a point, write: "The provided sources do not address this point."
- **Every factual claim must be cited** immediately after the claim using the reference number in DOUBLE brackets, e.g. [[1]] or [[3]]. To cite several sources at once, group them: [[1, 3]]. Always use double brackets so your citations are never confused with bracketed numbers that appear inside quoted source text.
- **When you quote a specific passage** and that source's header shows a page (e.g. "(p. 45)"), cite it as [[N, p. 45]]. If no page is shown, just use [[N]]. Do not invent page numbers.
- Write at least {target_words} words total. Develop each section fully using evidence from the excerpts.

## REPORT STRUCTURE

### 1. Executive Summary (300–500 words)
A concise overview of the topic drawing from the sources. Mention key themes, major concepts, and main findings. This should give the reader a complete but brief understanding of the topic.

### 2. Full Report
Develop the topic in depth with these sections:

1. **Introduction & Definition** — what do the sources say this concept is?
2. **Theoretical Foundations** — how do the sources describe its origins and place in Bowen theory?
3. **Key Dimensions** — what distinct aspects or components do the sources identify?
4. **Relationship to Other Bowen Concepts** — what connections do the sources explicitly draw?
5. **Clinical Presentation** — how do the sources describe this appearing in families or individuals?
6. **Clinical Implications & Therapeutic Approach** — what do the sources say about working with this clinically?
7. **Direct Quotations & Illustrations** — include key verbatim or near-verbatim passages from the sources
8. **Gaps & Limitations** — what does this topic lack coverage on in the provided sources?

**IMPORTANT: Do NOT include a References section in your output.** The reference list will be appended automatically. Just use [[1]], [[2]], etc. (double brackets) for inline citations throughout your report, using the numbers from the source list below.

## Source numbers (for inline citations only — do NOT reproduce this list in your output)
{refs_md}
"""

        st.subheader("Report")
        system = st.session_state.get("system_prompt", SYSTEM_PROMPT)
        style  = citations.normalize_style(
            st.session_state.get("citation_style", citations.DEFAULT_STYLE))
        report_ph = st.empty()
        try:
            acc: list = []
            for tok in _llm_stream([{"role": "user", "content": prompt}], system):
                acc.append(tok)
                if len(acc) % 12 == 0:          # throttle live re-render
                    report_ph.markdown("".join(acc))
            result = "".join(acc)
        except Exception as e:
            st.error(f"LLM error: {e}")
            return

        # Rewrite [[N]] markers into the chosen style and build the reference list from
        # ONLY the sources actually cited. Guard the whole post-process: a human-edited
        # sources.yml record must never crash the page and lose the streamed report.
        try:
            styled_body = citations.apply_intext_citations(result, num_to_record, style)
            raw_cited = citations.cited_numbers(result, set(num_to_record))
            cited = raw_cited or set(num_to_record)
            refs_body = citations.build_reference_list_md(num_to_record, style, cited)
            final_report = styled_body + f"\n\n## References\n\n{refs_body}\n"
            if not raw_cited and len(result.strip()) > 200:
                # Non-empty report but zero [[N]] markers → the model ignored the required
                # citation format. Surface it loudly instead of silently listing all sources.
                note, warn = ("No [[N]] citation markers were found — the model may not have "
                              "used the required format, so in-text citations are unstyled and "
                              "the reference list shows all retrieved sources. Try regenerating."), True
            else:
                verified_n = sum(1 for n in cited if num_to_record[n].get("verified"))
                note, warn = (f"Citations in {style} style · {verified_n}/{len(cited)} cited "
                              "sources have verified bibliographic data (edit sources.yml).", False)
        except Exception as e:
            plain = re.sub(r'\[\[\s*(\d[^\]]*?)\s*\]\]', r'[\1]', result)
            plain_refs = "\n".join(f"{n}. {nm}"
                                   for nm, n in sorted(ref_map.items(), key=lambda x: x[1]))
            final_report = plain + f"\n\n## References\n\n{plain_refs}\n"
            note, warn = f"Citation styling failed ({e}); showing plain numbered references.", True
        report_ph.markdown(final_report)                 # final view replaces raw stream
        (st.warning if warn else st.caption)(note)
        st.session_state.last_report = final_report

        # Build appendix from source texts
        if include_appendix and docs:
            appendix_parts = ["\n\n---\n\n## Appendix: Source Texts\n"]
            for doc_name in sorted(docs, key=lambda d: ref_map[d]):
                appendix_parts.append(f"\n### [{ref_map[doc_name]}] {doc_name}\n")
                for txt in docs[doc_name]:
                    appendix_parts.append(_format_chunk_text(txt) + "\n")
                appendix_parts.append("\n---\n")
            st.session_state.last_rpt_appendix = "\n".join(appendix_parts)
        else:
            st.session_state.last_rpt_appendix = ""

    # Show chunks used (after generation)
    if st.session_state.get("last_rpt_context"):
        with st.expander("Audit: show chunks sent to LLM"):
            sections = re.split(r'\n\n---\n\n', st.session_state.last_rpt_context)
            for section in sections:
                lines     = section.strip().split('\n', 1)
                hdr       = lines[0].lstrip('#').strip()
                body      = re.sub(r'\n{3,}', '\n\n', lines[1].strip()) if len(lines) > 1 else ""
                st.markdown(f"**{hdr}**")
                st.text(body)
                st.divider()

    # Appendix (shown after audit, before download)
    if st.session_state.get("last_rpt_appendix"):
        with st.expander("Appendix: Source Texts"):
            st.markdown(st.session_state.last_rpt_appendix)

    # Download buttons — includes appendix when present
    if st.session_state.get("last_report"):
        full_download = st.session_state.last_report
        if st.session_state.get("last_rpt_appendix"):
            full_download += "\n\n" + st.session_state.last_rpt_appendix

        # Build a filename stem from the topic
        topic_stub = (st.session_state.get("last_search_query", "bowen_report")
                      .strip()[:50].replace(" ", "_").replace("/", "-").replace(":", ""))
        if not topic_stub:
            topic_stub = "bowen_report"

        st.markdown("**Download Report**")
        dl1, dl2, dl3 = st.columns(3)
        with dl1:
            st.download_button(
                "📄 Markdown (.md)",
                data=full_download,
                file_name=f"{topic_stub}.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with dl2:
            try:
                docx_bytes = md_to_docx_bytes(full_download)
                st.download_button(
                    "📝 Word (.docx)",
                    data=docx_bytes,
                    file_name=f"{topic_stub}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"Word export unavailable: {e}")
        with dl3:
            try:
                pdf_bytes = md_to_pdf_bytes(full_download)
                st.download_button(
                    "📕 PDF (.pdf)",
                    data=pdf_bytes,
                    file_name=f"{topic_stub}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"PDF export unavailable: {e}")


def page_settings():
    st.header("Settings")

    # ── Search defaults ──────────────────────────────────────────────────────
    st.subheader("Search Defaults")

    all_modes = [
        ("Hybrid (BM25 + Embedding) — recommended", "hybrid"),
        ("Top Docs",                                 "top-docs"),
        ("Semantic (TF-IDF)",                        "semantic"),
        ("Keyword",                                  "keyword"),
        ("Both (Semantic + Keyword)",                "both"),
        ("Embedding",                                "embedding"),
    ]
    all_mode_labels = [m[0] for m in all_modes]
    all_mode_values = [m[1] for m in all_modes]
    cur_default     = st.session_state.get("default_search_mode", "hybrid")
    cur_idx         = all_mode_values.index(cur_default) if cur_default in all_mode_values else 0

    chosen = st.selectbox(
        "Default search mode",
        range(len(all_mode_labels)),
        format_func=lambda i: all_mode_labels[i],
        index=cur_idx,
        help="Pre-selected mode on the Search, Report, and Chat pages when you first load them. "
             "Hybrid and Embedding require the embedding index to be built.",
    )
    st.session_state.default_search_mode = all_mode_values[chosen]

    if all_mode_values[chosen] in ("hybrid", "embedding") and not (
        EMBEDDING_AVAILABLE and getattr(_get_index(), "embed_matrix", None) is not None
    ):
        st.warning("Hybrid and Embedding modes require the embedding index. "
                   "If it is not available the pages will fall back to Top Docs.")

    st.divider()
    st.subheader("Citations")
    _cur_style = citations.normalize_style(
        st.session_state.get("citation_style", citations.DEFAULT_STYLE))
    _style = st.selectbox(
        "Citation style (Report references + in-text citations)",
        citations.STYLES,
        index=citations.STYLES.index(_cur_style),
        help="How the Report lists its references and formats in-text citations. "
             "APA/MLA/Chicago/Harvard use author–date or author–page; Vancouver is "
             "numbered [1]. Full citations read from sources.yml; missing fields show "
             "honestly (e.g. 'n.d.') and are never invented.",
    )
    st.session_state.citation_style = _style
    _verified = sum(1 for r in SOURCES if r.get("verified"))
    if SOURCES:
        st.caption(f"sources.yml: {len(SOURCES)} records, {_verified} verified. "
                   "Unverified records still cite with best-available data; verify the "
                   "works you cite (author, year, title, publisher) for exact references.")
    else:
        st.caption("No sources.yml found — references fall back to document filenames. "
                   "Run `python3 seed_sources.py` to create one.")

    st.divider()
    st.subheader("LLM Provider")
    _providers = ["claude", "openai", "deepseek", "ollama"]
    _cur_provider = st.session_state.get("provider", "deepseek")
    if _cur_provider not in _providers:
        _cur_provider = "deepseek"
    provider = st.radio("Provider", _providers,
                        index=_providers.index(_cur_provider),
                        horizontal=True)
    st.session_state.provider = provider

    def _key_input(label: str, ss_key: str):
        existing = st.session_state.get(ss_key, "")
        if existing:
            st.caption(f"Key saved: ···{existing[-6:]}")
        new_val = st.text_input(
            label, value="", type="password",
            placeholder="Paste new key to update…" if existing else "Paste key here…",
        )
        if new_val:
            st.session_state[ss_key] = new_val

    if provider == "claude":
        st.subheader("Claude (Anthropic)")
        _key_input("API Key", "claude_key")
        model = st.selectbox("Model", CLAUDE_MODELS,
                             index=CLAUDE_MODELS.index(
                                 st.session_state.get("claude_model", "claude-sonnet-4-6"))
                             if st.session_state.get("claude_model") in CLAUDE_MODELS else 0)
        st.session_state.claude_model = model

    elif provider == "openai":
        st.subheader("OpenAI")
        _key_input("API Key", "openai_key")
        model = st.selectbox("Model", OPENAI_MODELS,
                             index=OPENAI_MODELS.index(
                                 st.session_state.get("openai_model", "gpt-4o"))
                             if st.session_state.get("openai_model") in OPENAI_MODELS else 0)
        st.session_state.openai_model = model

    elif provider == "deepseek":
        st.subheader("DeepSeek")
        _key_input("API Key", "deepseek_key")
        model = st.selectbox("Model", DEEPSEEK_MODELS,
                             index=DEEPSEEK_MODELS.index(
                                 st.session_state.get("deepseek_model", "deepseek-v4-flash"))
                             if st.session_state.get("deepseek_model") in DEEPSEEK_MODELS else 0)
        st.session_state.deepseek_model = model
        st.caption(f"Endpoint: {DEEPSEEK_BASE_URL}")

    else:
        st.subheader("Ollama (self-hosted)")
        st.info("Ollama must be running and accessible from the server. "
                "On Railway this requires a separately hosted Ollama instance.")
        url = st.text_input("Server URL", value=st.session_state.get("ollama_url",
                                                                       "http://localhost:11434"))
        st.session_state.ollama_url = url
        model = st.text_input("Model", value=st.session_state.get("ollama_model", "qwen2.5:7b"))
        st.session_state.ollama_model = model

    st.divider()
    st.subheader("System Prompt")
    sp = st.text_area("System Prompt", value=st.session_state.get("system_prompt", SYSTEM_PROMPT),
                      height=200)
    st.session_state.system_prompt = sp

    st.divider()
    if st.button("Test connection", type="primary"):
        with st.spinner("Testing…"):
            try:
                result = "".join(_llm_stream(
                    [{"role": "user", "content": "Reply with exactly: OK"}],
                    "You are a test assistant."
                ))
                st.success(f"Connected — response: {result[:80]}")
            except Exception as e:
                st.error(f"Connection failed: {e}")


def page_index(idx: IndexManager):
    st.header("Index")

    if idx.loaded:
        docs  = len(set(c["doc_name"] for c in idx.chunks))
        st.metric("Documents", docs)
        st.metric("Chunks", f"{len(idx.chunks):,}")
        st.metric("Embeddings", "loaded" if idx.embed_matrix is not None else "not available")
        st.metric("BM25", "loaded" if idx.bm25 is not None else "not available")
    else:
        st.error("Index not loaded.")

    st.divider()
    st.info(
        "**To update the index:** rebuild locally with the desktop app "
        "(Index tab → Rebuild Index), then commit and push the updated "
        "`rag-document-search/references/` files. Railway redeploys automatically."
    )

    with st.expander("Document list"):
        for doc in idx.list_documents():
            author = doc_author(doc)
            boost  = authority_boost(doc)
            badge  = f" ★ {boost}×" if boost > 1.0 else ""
            st.caption(f"**{doc}**  —  {author}{badge}")


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="Bowen Theory RAG",
        layout="wide",
        initial_sidebar_state="expanded",
    )


    _init_session()
    _check_auth()

    idx = _get_index()

    _NAV = [
        ("Search",   "Find and browse source passages in the index. "
                     "Check results and stage them to carry into a Report."),
        ("Chat",     "Conversational Q&A with the literature. "
                     "Ask follow-up questions; each turn retrieves fresh source chunks."),
        ("Report",   "Generate a structured, cited research report on a topic. "
                     "Best for getting a comprehensive, referenced answer. "
                     "This is the core workflow for most users."),
        ("Index",    "Admin — view index statistics (document count, chunk count, embedding status). "
                     "Do not modify unless you are rebuilding the index."),
        ("Settings", "Admin — configure the LLM provider, API keys, default search mode, "
                     "and system prompt. Do not change unless you know what you are doing."),
        ("User Guide", "Full documentation — search modes, score badges, Chat, Report, "
                       "and admin workflows for index management and deployment."),
    ]

    if "nav_page" not in st.session_state:
        st.session_state["nav_page"] = "Search"

    with st.sidebar:
        st.markdown("### Bowen Theory RAG")
        _provider = st.session_state.get("provider", "claude")
        _model = {
            "claude":   st.session_state.get("claude_model",   "claude-sonnet-4-6"),
            "openai":   st.session_state.get("openai_model",   "gpt-4o"),
            "deepseek": st.session_state.get("deepseek_model", "deepseek-v4-flash"),
            "ollama":   st.session_state.get("ollama_model",   "qwen2.5:7b"),
        }.get(_provider, _provider)
        if idx.loaded:
            docs  = len(set(c["doc_name"] for c in idx.chunks))
            embed = f" · {len(idx.embed_matrix):,} embeddings" if idx.embed_matrix is not None else ""
            st.markdown(
                f'<div style="font-size:12px;color:#4b5563;line-height:1.4;margin:0">'
                f'{docs} docs · {len(idx.chunks):,} chunks{embed}<br>'
                f'🤖 {_provider} · {_model}</div>',
                unsafe_allow_html=True,
            )
        st.divider()
        for _name, _desc in _NAV:
            _selected = st.session_state["nav_page"] == _name
            _nc, _hc = st.columns([5, 1])
            with _nc:
                if st.button(
                    ("▶ " if _selected else "") + _name,
                    use_container_width=True,
                    key=f"nav_{_name}",
                    type="primary" if _selected else "secondary",
                ):
                    st.session_state["nav_page"] = _name
                    st.session_state[f"show_help_{_name}"] = False
                    st.rerun()
            with _hc:
                if st.button("?", key=f"help_{_name}", use_container_width=True):
                    _tog = f"show_help_{_name}"
                    st.session_state[_tog] = not st.session_state.get(_tog, False)
                    st.rerun()
            if st.session_state.get(f"show_help_{_name}", False):
                st.info(_desc)
        st.divider()
        st.caption("Bowen Family Systems Theory research tool")

    page = st.session_state["nav_page"]

    if page == "Search":
        page_search(idx)
    elif page == "Chat":
        page_chat(idx)
    elif page == "Report":
        page_report(idx)
    elif page == "Index":
        page_index(idx)
    elif page == "Settings":
        page_settings()
    else:
        page_user_guide()


def page_user_guide():
    guide_path = BASE_DIR / "USER_GUIDE.md"
    if guide_path.exists():
        st.markdown(guide_path.read_text(encoding="utf-8"))
    else:
        st.error("USER_GUIDE.md not found. It should be in the project root directory.")


if __name__ == "__main__":
    main()
