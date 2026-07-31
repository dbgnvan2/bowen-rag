"""
Export round-trip test (M4.C).

A citation-styled report (author-date in-text + a References section) must convert
to .docx and .pdf without error and preserve the styled citations. Exercises the
REAL converters in streamlit_app (md_to_docx_bytes / md_to_pdf_bytes).

Run: python3 test_report_export.py
Note: imports streamlit_app, which pulls in the ML stack — slow but real.
"""

import io
import unittest

import citations as C
import streamlit_app as app

BOOK = {"type": "book", "authors": [{"family": "Bowen", "given": "Murray"}],
        "year": 1978, "title": "Family Therapy in Clinical Practice",
        "publisher": "Jason Aronson"}


class TestExportWithCitations(unittest.TestCase):
    def setUp(self):
        num_to_record = {1: BOOK}
        body = ("## Executive Summary\n\n"
                "Differentiation is central to the theory [[1]]. "
                'A direct quote: "the family is a system" [[1, p. 45]].\n\n'
                "### Key Dimensions\n\n- Emotional systems shape behaviour [[1]]\n")
        styled = C.apply_intext_citations(body, num_to_record, "APA")
        refs = C.build_reference_list_md(num_to_record, "APA")
        self.report = styled + f"\n\n## References\n\n{refs}\n"
        # Sanity: the fixture really is styled (no numbered marker left in body).
        self.assertIn("(Bowen, 1978)", self.report)
        self.assertIn("(Bowen, 1978, p. 45)", self.report)

    def test_m4c_docx_export_preserves_citations(self):
        data = app.md_to_docx_bytes(self.report, title="Test Report")
        self.assertIsInstance(data, (bytes, bytearray))
        self.assertGreater(len(data), 0)
        from docx import Document
        doc = Document(io.BytesIO(data))
        full = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("(Bowen, 1978)", full)         # in-text citation survived
        self.assertIn("p. 45", full)                 # quote page locator survived
        self.assertIn("References", full)            # references heading present
        self.assertNotIn("[1]", full)                # numbered marker was rewritten

    def test_m4c_pdf_export_valid(self):
        data = app.md_to_pdf_bytes(self.report, title="Test Report")
        self.assertTrue(data.startswith(b"%PDF"))    # valid PDF header
        self.assertGreater(len(data), 800)

    def test_m4c_vancouver_numbered_export(self):
        num_to_record = {1: BOOK}
        body = "Claim [[1]]. Another [[1]]."
        styled = C.apply_intext_citations(body, num_to_record, "Vancouver")
        refs = C.build_reference_list_md(num_to_record, "Vancouver")
        report = styled + f"\n\n## References\n\n{refs}\n"
        data = app.md_to_docx_bytes(report, title="V")
        from docx import Document
        full = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        self.assertIn("[1]", full)                   # numbered style keeps markers
        self.assertIn("Bowen M.", full)              # Vancouver reference entry


if __name__ == "__main__":
    unittest.main(verbosity=2)
